import sys
import matplotlib.pyplot as plt
import numpy as np
import docx
from matplotlib import pyplot
from mpl_toolkits.mplot3d import Axes3D


# top-1
# model_name 0.6, 0.65, 0.7, 0.75, 0.8, 0.85, 0.9, 0.95


# top-5
# model_name 0.6, 0.65, 0.7, 0.75, 0.8, 0.85, ...



# dic[quant_bits] = dic[model_name] = dic[0.9] = [top-1, top5]


# 3d preview  ....  model x 2(top-1, top-5), {x:prune, y:sparsity, z:acc}

file_name = sys.argv[1]

quant_bits_dic = {}

def li(l, fd, format = False):
    if format == True:
        l = [":-:" for i in range(len(l))]
    print("|", file = fd, end = "")
    for elem in l:
        print(str(elem) + "|", file = fd, end = "")
    print("", file = fd)


def dfs(dic):
    if "top-1" in dic.keys():
        print(dic["top-1"])
        print(dic["top-5"])
        return
    for key in dic.keys():
        print(key)
        dfs(dic[key])


def get_dic(l):
    dic = {}
    key = None
    value = None
    for index, elem in enumerate(l):
        if elem == '=':
            dic[l[index - 1].split(',')[0]] = l[index + 1].split(',')[0]
    return dic


def handel(file_name):
    with open(file_name, "r") as fd:
        text = fd.readlines()
        state = 0
        quant_bits = "False"
        for sentence in text:
            print(sentence, end = "")
            word = sentence.split()
            print(word)
            if "quant_bits" in word:
                dic = get_dic(word)
                print(dic)
                quant_bits = dic["quant_bits"]
                model_name = dic["model_name"]
                enable_prune = dic["enable_prune"]
                if quant_bits not in quant_bits_dic:
                    quant_bits_dic[quant_bits] = {}
                if model_name not in quant_bits_dic[quant_bits]:
                    quant_bits_dic[quant_bits][model_name] = {}
                if enable_prune not in quant_bits_dic[quant_bits][model_name]:
                    quant_bits_dic[quant_bits][model_name][enable_prune] = {"top-1":[], "top-5": []}
            if "top_1_acc" in word:
                dic = get_dic(word)
                quant_bits_dic[quant_bits][model_name][enable_prune]["top-1"].append(dic["top_1_acc"])
                quant_bits_dic[quant_bits][model_name][enable_prune]["top-5"].append(dic["top_5_acc"])



def row(model_name, dic, mark, choose, fd):
    line = [model_name]
    for key in dic.keys():
        try:
            line.append(dic[key][mark][choose])
        except IndexError:
            line.append(None)
    li(line, fd)


def table(dic, mark, fd):
    title = [str(float(i) / 100) for i in range(60, 100, 5)]
    title = ["model_name"] + title
    li(title, fd)
    li(title, fd, True)
    for model_name in dic.keys():
        row(model_name, dic[model_name], mark, 0, fd)
        row(model_name + "(pruned)", dic[model_name], mark, 1, fd)




def markdown(dic):
    with open("draft.md", "w") as fd:
        for key in dic.keys():
            print("## quant_bits = {}".format(key), file = fd, end = "\n\n")
            print("### top-1", end = "\n\n", file = fd)
            table(dic[key], "top-1", fd)
            print("### top-5", end = "\n\n", file = fd)
            table(dic[key], "top-5", fd)


def rgb_num(enable_inc = True, enable_one = False, enable_zero = False):
    def indentify(num):
        return num
    def re_indentify(num):
        return 1 - num
    def one(num):
        return 1
    def zero(num):
        return 0
    if enable_one:
        return one
    elif enable_zero:
        return zero
    if enable_inc:
        return indentify
    else:
        return re_indentify

rgb = [
        (rgb_num(), rgb_num(), rgb_num()),
        (rgb_num(enable_one = True), rgb_num(), rgb_num(enable_zero = True)), 
        (rgb_num(False), rgb_num(enable_one = True), rgb_num(enable_zero = True)),
        (rgb_num(enable_zero = True), rgb_num(enable_one = True), rgb_num()),
        (rgb_num(enable_zero = True), rgb_num(False), rgb_num(enable_one = True)), 
        (rgb_num(), rgb_num(enable_zero = True), rgb_num(enable_one = True))]

def data_flatten(dic):
    # dic[model_name + acc] = []
    model_name = ["Alexnet", "Vgg16", "Resnet18", "Resnet34"]
    quant_bits = ["4", "5", "6", "7", "8", "False"]
    sparsity   = ["0.6", "0.65", "0.7", "0.75", "0.8", "0.85", "0.9", "0.95"]
    sparsity.reverse()
    ret_dic = {}
    not_quant_ret_dic = {}
    for name in model_name:
        top_1 = []
        top_1_pruned = []
        top_5 = []
        top_5_pruned = []
        top_1_not_quant = []
        top_5_not_quant = []
        for bits in quant_bits:
            for s in sparsity:
                top_1.append(float(dic[bits][name][s]["top-1"][0]))
                top_1_pruned.append(float(dic[bits][name][s]["top-1"][1]))
                top_5.append(float(dic[bits][name][s]["top-5"][0]))
                top_5_pruned.append(float(dic[bits][name][s]["top-5"][1]))
                if bits == "False":
                    top_1_not_quant.append(float(dic[bits][name][s]["top-1"][1]))
                    top_5_not_quant.append(float(dic[bits][name][s]["top-5"][1]))
        ret_dic[name + "_top-1"] = top_1;
        ret_dic[name + "_top-1" + "_pruned"] = top_1_pruned;
        ret_dic[name + "_top-5"] = top_5;
        ret_dic[name + "_top-5" + "_pruned"] = top_5_pruned;
        top_1_not_quant.reverse()
        top_5_not_quant.reverse()
        not_quant_ret_dic[name + "_top-1" + "_pruned"] = top_1_not_quant
        not_quant_ret_dic[name + "_top-5" + "_pruned"] = top_5_not_quant
    for key in not_quant_ret_dic.keys():
        print(key, not_quant_ret_dic[key])
    return ret_dic, not_quant_ret_dic

            
def not_quant_plot(dic):
    def wrap(string, num):
        return string + "_top-" + str(num) + "_pruned"
    _, data_dic = data_flatten(dic)
    x = [str(i / 100) for i in range(60, 100, 5)]

    def sub_plot(pos, num):
        plt.subplot(pos)
        plt.plot(x, data_dic[wrap("Vgg16", num)], "bo-", label = "vgg16")
        plt.plot(x, data_dic[wrap("Resnet18", num)], 'y*-', label = "resnet18")
        plt.plot(x, data_dic[wrap("Resnet34", num)], 'r+-', label = "resnet34")
        plt.plot(x, data_dic[wrap("Alexnet", num)], 'g^-', label = "alexnet")
        plt.legend()
        plt.xlabel("train method")
        plt.ylabel("Top-" + str(num) + " accuracy")
    sub_plot(121, 1)
    sub_plot(122, 5)
    plt.show()



def plot(dic):
    data_dic, _= data_flatten(dic)
    x_quant_flatten = [i for i in range(6)]
    y_puren_flatten = [i for i in range(8)]
    xpos, ypos = np.meshgrid(x_quant_flatten, y_puren_flatten)
    xpos = xpos.flatten('F')
    ypos = ypos.flatten('F')
    print(xpos, ypos)
    zpos = np.zeros_like(xpos)

    dx = 0.3 * np.ones_like(zpos)
    dy = dx.copy()
    dz = np.array([i % 8 for i in range(6 * 8)])
    print(dx.shape, dy.shape, dz.shape)

    fig = plt.figure()
    c = []
    for j in range(6):
        for i in range(8):
            c.append((
                rgb[j][0](dz[j * 8 + i] / 8), 
                rgb[j][1](dz[j * 8 + i] / 8), 
                rgb[j][2](dz[j * 8 + i] / 8)))
    
    def plot_figure(fig, pos, name):
        ax = fig.add_subplot(pos, projection = '3d')
        # print(ax.spines)
        # print(ax.zlim)
        ax.set_xticklabels(['fp','8bit','7bit','6bit','5bit','4bit'])
        ax.set_yticklabels([str(float(i) / 100) for i in range(95, 55, -5)])
        ax.set_xlabel("quant_bits")
        ax.set_ylabel("sparsity ratio")
        ax.set_zlabel(name.split('_')[1] + " accuracy")
        ax.set_title(name)
        ax.bar3d(xpos, ypos, zpos, dx, dy, data_dic[name], color=c, zsort='average')

    # plot_figure(fig, 221, "Alexnet_top-1_pruned")
    # plot_figure(fig, 222, "Vgg16_top-1_pruned")
    # plot_figure(fig, 223, "Alexnet_top-5_pruned")
    # plot_figure(fig, 224, "Vgg16_top-5_pruned")
    

    plot_figure(fig, 221, "Resnet18_top-1_pruned")
    plot_figure(fig, 222, "Resnet34_top-1_pruned")
    plot_figure(fig, 223, "Resnet18_top-5_pruned")
    plot_figure(fig, 224, "Resnet34_top-5_pruned")
    
    
    plt.show()
    # plt.savefig("test.png", figsize = (160, 90), dpi = 900)

def one_table(dic, doc, top):
    def change_num(string):
        string = string.split('.')[1]
        return string[0:2] + "." + string[2:4] + "%"
    table = doc.add_table(rows = 5, cols = 9, style = "Table Grid")
    model_name = ["None", "Alexnet", "Vgg16", "Resnet18", "Resnet34"]
    sparsity   = ["modelname", "0.6", "0.65", "0.7", "0.75", "0.8", "0.85", "0.9", "0.95"]
    for i, name in enumerate(model_name):
        for j, spar in enumerate(sparsity):
            if i == 0:
                table.rows[i].cells[j].text = sparsity[j]
            elif j == 0:
                table.rows[i].cells[j].text = model_name[i]
            else:
                table.rows[i].cells[j].text = change_num(dic[name][spar][top][1])
                

        
    

def make_docx(dic):
    document = docx.Document()
    quant_bits = ["4", "5", "6", "7", "8", "False"]
    for bits in quant_bits:
        document.add_paragraph("quant_bits = {}, top-1".format(bits))
        one_table(dic[bits], document, "top-1")
        document.add_paragraph("quant_bits = {}, top-5".format(bits))
        one_table(dic[bits], document, "top-5")
    document.save("list.docx")


def main():
    for i in range(1, len(sys.argv)):
        handel(sys.argv[i])
    dfs(quant_bits_dic)
    # markdown(quant_bits_dic)
    # make_docx(quant_bits_dic)
    # plot(quant_bits_dic)
    not_quant_plot(quant_bits_dic)



main()

            
    
